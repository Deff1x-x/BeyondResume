from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from app.utils.resume_parse import (
    DOCX_MIME,
    EmptyExtractedResumeTextError,
    ResumeDocumentParseError,
    UnsupportedResumeParserTypeError,
    extract_resume_text,
    normalize_resume_text,
)


def test_docx_extracts_paragraphs_and_table_cells(tmp_path: Path) -> None:
    path = tmp_path / "resume.docx"
    document = Document()
    document.add_paragraph("Python Developer")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "FastAPI"
    document.add_paragraph("PostgreSQL")
    document.save(path)

    assert extract_resume_text(path, DOCX_MIME) == "Python Developer\nFastAPI\nPostgreSQL"


def write_pdf(path: Path, pages: list[str]) -> None:
    objects = ["<< /Type /Catalog /Pages 2 0 R >>", None]
    page_ids: list[int] = []
    for text in pages:
        page_ids.append(len(objects) + 1)
        objects.extend(
            [
                "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 0 0 R >> >> /Contents 0 0 R >>",
                "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
                f"<< /Length {len(f'BT /F1 12 Tf 72 720 Td ({text}) Tj ET')} >>\nstream\nBT /F1 12 Tf 72 720 Td ({text}) Tj ET\nendstream",
            ]
        )
        objects[-3] = (
            objects[-3]
            .replace("0 0 R", f"{len(objects) - 1} 0 R", 1)
            .replace("0 0 R", f"{len(objects)} 0 R", 1)
        )
    objects[1] = (
        f"<< /Type /Pages /Kids [{' '.join(f'{page_id} 0 R' for page_id in page_ids)}] /Count {len(pages)} >>"
    )
    output = "%PDF-1.4\n"
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(output.encode()))
        output += f"{index} 0 obj\n{obj}\nendobj\n"
    xref = len(output.encode())
    output += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n"
    output += "".join(f"{offset:010} 00000 n \n" for offset in offsets[1:])
    output += f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF\n"
    path.write_bytes(output.encode())


def test_pdf_extracts_text_and_preserves_page_order(tmp_path: Path) -> None:
    from app.utils.resume_parse import PDF_MIME

    path = tmp_path / "resume.pdf"
    write_pdf(path, ["First Page", "Second Page"])

    assert extract_resume_text(path, PDF_MIME) == "First Page\nSecond Page"


def test_pdf_none_page_text_is_empty(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    from app.utils import resume_parse
    from app.utils.resume_parse import PDF_MIME

    path = tmp_path / "resume.pdf"
    path.write_bytes(b"placeholder")

    class Page:
        def extract_text(self):
            return None

    class Reader:
        is_encrypted = False
        pages = [Page()]

    monkeypatch.setattr(resume_parse, "PdfReader", lambda _path: Reader())
    with pytest.raises(EmptyExtractedResumeTextError):
        extract_resume_text(path, PDF_MIME)


def test_corrupted_and_encrypted_pdf_fail(tmp_path: Path) -> None:
    from app.utils.resume_parse import PDF_MIME

    corrupted = tmp_path / "broken.pdf"
    corrupted.write_bytes(b"not a pdf")
    with pytest.raises(ResumeDocumentParseError):
        extract_resume_text(corrupted, PDF_MIME)

    encrypted = tmp_path / "encrypted.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=100, height=100)
    writer.encrypt("secret")
    with encrypted.open("wb") as output:
        writer.write(output)
    with pytest.raises(ResumeDocumentParseError):
        extract_resume_text(encrypted, PDF_MIME)


def test_normalization_is_deterministic_and_preserves_unicode() -> None:
    assert (
        normalize_resume_text(" \u0410\u043b\u0430\u043d\t  DEV\r\n\r\n\rtext  ")
        == "\u0410\u043b\u0430\u043d DEV\n\ntext"
    )


def test_unsupported_and_invalid_documents_raise_typed_errors(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a document")

    with pytest.raises(UnsupportedResumeParserTypeError):
        extract_resume_text(path, "text/plain")
    with pytest.raises(ResumeDocumentParseError):
        extract_resume_text(path, DOCX_MIME)


def test_empty_docx_is_not_successfully_parsed(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    Document().save(path)

    with pytest.raises(EmptyExtractedResumeTextError):
        extract_resume_text(path, DOCX_MIME)
